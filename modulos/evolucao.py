import streamlit as st
from datetime import date, datetime
import matplotlib.pyplot as plt
from collections import Counter
from docx import Document
import os
from services import listar_clientes, listar_atendimentos

def render(cursor, conn):
    st.header("📈 Evolução por Cliente")

    termo_busca = st.text_input("Digite o nome do cliente")
    if termo_busca:
        cursor.execute("SELECT id, nome FROM clientes WHERE nome LIKE ?", (f"%{termo_busca}%",))
        resultados = cursor.fetchall()

        if len(resultados) == 0:
            st.warning("Nenhum cliente encontrado.")
            return
        elif len(resultados) == 1:
            cliente_id, cliente_nome = resultados[0]
        else:
            cliente_nome = st.selectbox("Selecione o cliente:", [nome for _, nome in resultados])
            cliente_id = next(cid for cid, nome in resultados if nome == cliente_nome)

        cursor.execute("""
            SELECT data, prontuario, categorias
            FROM atendimentos
            WHERE cliente_id = ?
            ORDER BY data ASC
        """, (cliente_id,))
        historico = cursor.fetchall()

        if historico:
            st.subheader(f"📝 Histórico de atendimentos - {cliente_nome}")
            for data, prontuario, categorias in historico:
                st.markdown(f"**📅 {data}**")
                if categorias:
                    st.markdown(f"**🏷️ Categorias:** {categorias}")
                st.markdown(f"📝 {prontuario}")
                st.markdown("---")

            # Categorias mais frequentes
            todas_categorias = []
            for _, _, categorias in historico:
                if categorias:
                    todas_categorias.extend([c.strip() for c in categorias.split(",")])
            contagem = Counter(todas_categorias)

            if contagem:
                st.subheader("🏷️ Categorias mais frequentes")
                fig, ax = plt.subplots()
                ax.bar(contagem.keys(), contagem.values())
                ax.set_ylabel("Frequência")
                ax.set_xticklabels(contagem.keys(), rotation=45, ha='right')
                st.pyplot(fig)

            # Frequência média de sessões
            datas = [registro[0] for registro in historico]
            if len(datas) > 1:
                datas_convertidas = [datetime.strptime(d, "%Y-%m-%d") for d in datas]
                datas_convertidas.sort()
                difs = [(datas_convertidas[i+1] - datas_convertidas[i]).days for i in range(len(datas_convertidas)-1)]
                media_dias = sum(difs) / len(difs)
                st.subheader("📆 Frequência média de sessões")
                if media_dias <= 10:
                    frequencia = "Semanal"
                elif media_dias <= 20:
                    frequencia = "Quinzenal"
                else:
                    frequencia = "Esporádica"
                st.info(f"Frequência estimada: **{frequencia}** ({round(media_dias)} dias em média entre sessões)")

            # Exportar para Word
            if st.button("📄 Exportar para Word"):
                doc = Document()
                doc.add_heading(f"Relatório de Atendimento - {cliente_nome}", 0)
                for i, (data, prontuario, categorias) in enumerate(historico, 1):
                    doc.add_heading(f"Atendimento {i}", level=1)
                    doc.add_paragraph(f"📅 Data: {data}")
                    doc.add_paragraph(f"🏷️ Categorias: {categorias if categorias else 'Nenhuma'}")
                    doc.add_paragraph("📝 Prontuário:")
                    doc.add_paragraph(prontuario if prontuario else "Sem anotações.")
                    doc.add_paragraph("\n")

                caminho = f"relatorio_{cliente_nome.replace(' ', '_')}.docx"
                doc.save(caminho)
                with open(caminho, "rb") as f:
                    st.download_button("📥 Baixar relatório", f, file_name=caminho,
                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                os.remove(caminho)

        else:
            st.info("Este cliente ainda não possui atendimentos registrados.")
