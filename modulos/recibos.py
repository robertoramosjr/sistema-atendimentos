import streamlit as st
from datetime import date, timedelta
from docx import Document
from docx.shared import Pt
import io
from services import valor_por_extenso

NOME_PSICOLOGA = "Ana Carolina de Sousa"
CPF_PSICOLOGA = "432.888.538-33"

def render(cursor, conn):
    st.subheader("📄 Gerar Recibo de Pagamento")

    cursor.execute("SELECT id, nome, cpf FROM clientes ORDER BY nome")
    clientes = cursor.fetchall()
    nomes_clientes = [nome for _, nome, _ in clientes]
    cliente_selecionado = st.selectbox("Selecione um cliente", nomes_clientes)
    cliente_id, _, cpf_cliente = next((cid, nome, cpf) for cid, nome, cpf in clientes if nome == cliente_selecionado)

    tipo_recibo = st.radio("Tipo de recibo", ["Mensal", "Anual"], horizontal=True)
    hoje = date.today()
    ano = st.number_input("Ano", min_value=2000, max_value=hoje.year, value=hoje.year)

    if tipo_recibo == "Mensal":
        mes = st.selectbox("Mês", list(range(1, 13)), format_func=lambda m: date(1900, m, 1).strftime('%B'))
        inicio = date(ano, mes, 1)
        fim = date(ano, mes, 28) + timedelta(days=4)
        fim = fim - timedelta(days=fim.day)
        titulo_periodo = f"Recibo referente a {inicio.strftime('%B/%Y')}"
    else:
        inicio = date(ano, 1, 1)
        fim = date(ano, 12, 31)
        titulo_periodo = f"Recibo referente ao ano de {ano}"

    cursor.execute("""
        SELECT data, valor, descricao FROM pagamentos
        WHERE cliente_id = ? AND data BETWEEN ? AND ?
        ORDER BY data
    """, (cliente_id, inicio.isoformat(), fim.isoformat()))
    pagamentos = cursor.fetchall()

    if pagamentos:
        if st.button("📄 Gerar recibo"):
            doc = Document()
            estilo = doc.styles['Normal']
            estilo.font.name = 'Arial'
            estilo.font.size = Pt(11)

            doc.add_heading("Recibo de Pagamento", level=1)
            doc.add_paragraph(titulo_periodo)
            doc.add_paragraph(f"Nome do cliente: {cliente_selecionado}")
            if cpf_cliente:
                doc.add_paragraph(f"CPF do cliente: {cpf_cliente}")

            doc.add_paragraph("")
            table = doc.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Data'
            hdr_cells[1].text = 'Valor'
            hdr_cells[2].text = 'Descrição'

            total = 0
            for data_pgto, valor, descricao in pagamentos:
                row_cells = table.add_row().cells
                row_cells[0].text = date.fromisoformat(data_pgto).strftime('%d/%m/%Y')
                row_cells[1].text = f"R$ {valor:.2f}"
                row_cells[2].text = descricao
                total += valor

            doc.add_paragraph("")
            doc.add_paragraph(
                f"Eu, {NOME_PSICOLOGA}, CPF: {CPF_PSICOLOGA}, recebi de {cliente_selecionado}"
                f" o valor de R$ {total:.2f} ({valor_por_extenso(total)} reais), referente a atendimentos psicológicos."
            )

            doc.add_paragraph("\n\nCidade, ____/____/______")
            doc.add_paragraph("\nAssinatura: _________________________________")

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            nome_arquivo = f"recibo_{cliente_selecionado.lower().replace(' ', '_')}_{inicio.strftime('%Y%m') if tipo_recibo == 'Mensal' else ano}.docx"
            st.download_button(
                label="📥 Baixar recibo (.docx)",
                data=buffer,
                file_name=nome_arquivo,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.info("Nenhum pagamento registrado para esse período.")
        